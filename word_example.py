import docx

document = docx.Document('demo.docx')
print(document.paragraphs[0].text)

paragraph = document.paragraphs[1]
run = paragraph.runs[2]
print(run.text)

run.italic = True
run.underline = True
paragraph.add_run('This is a new run.')

document.add_paragraph('Hello this is a new paragraph.')

document.save('demo2.docx')